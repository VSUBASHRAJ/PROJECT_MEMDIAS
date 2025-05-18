import os
import logging
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, session
from flask_login import LoginManager, UserMixin, current_user, login_user, logout_user, login_required
import pyrebase
import uuid
import base64
import json
import tempfile
import io
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import openpyxl
from weasyprint import HTML, CSS
from flask import send_file


from dotenv import load_dotenv  # <-- Load this

# Load environment variables from .env file
load_dotenv()
# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "your-secret-key")

# Firebase Configuration
firebase_config = {
    "apiKey": os.getenv("FIREBASE_API_KEY"),
    "authDomain": os.getenv("FIREBASE_AUTH_DOMAIN"),
    "projectId": os.getenv("FIREBASE_PROJECT_ID"),
    "storageBucket": os.getenv("FIREBASE_STORAGE_BUCKET"),
    "messagingSenderId": os.getenv("FIREBASE_MESSAGING_SENDER_ID"),
    "appId": os.getenv("FIREBASE_APP_ID"),
    "databaseURL": os.getenv("FIREBASE_DATABASE_URL"),
}

# Initialize Firebase
firebase = pyrebase.initialize_app(firebase_config)
auth = firebase.auth()
db = firebase.database()
storage = firebase.storage()

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# User class for Flask-Login
class User(UserMixin):
    def __init__(self, id, email, token=None, display_name=None):
        self.id = id
        self.email = email
        self.token = token
        self.display_name = display_name
        
@login_manager.user_loader
def load_user(user_id):
    if 'user_email' in session:
        return User(user_id, session['user_email'], session.get('user_token'), session.get('display_name'))
    return None

# Routes
@app.route('/')
def index():
    """Render the main memories list page."""
    # If not logged in, still show the page but with a prompt to log in
    memories = []
    
    if current_user.is_authenticated:
        try:
            # Get search/filter parameters
            search = request.args.get('search', '')
            date_filter = request.args.get('date', '')
            time_from = request.args.get('time_from', '')
            time_to = request.args.get('time_to', '')
            
            # Get memories from Firebase
            memories_ref = db.child("memories").child(current_user.id).get(token=current_user.token)
            
            if memories_ref.val():
                memories = []
                for key, value in memories_ref.val().items():
                    memory = value
                    memory['id'] = key
                    
                    # Apply filters if provided
                    if search and search.lower() not in memory['title'].lower() and search.lower() not in memory['content'].lower():
                        continue
                    
                    if date_filter and memory['date'] != date_filter:
                        continue
                        
                    # Apply time range filter if provided
                    if time_from and memory['time'] < time_from:
                        continue
                        
                    if time_to and memory['time'] > time_to:
                        continue
                    
                    memories.append(memory)
        except Exception as e:
            logger.error(f"Error getting memories: {str(e)}")
            flash(f"Error loading memories: {str(e)}", "danger")
    
    return render_template('index.html', memories=memories)

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Handle login with email/password."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    error = None
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        try:
            # Sign in with email/password
            user = auth.sign_in_with_email_and_password(email, password)
            
            # Get user info and JWT token
            user_id = user['localId']
            user_email = user['email']
            user_token = user['idToken']
            
            # Check if user profile exists
            user_profile = db.child("users").child(user_id).get(token=user_token).val()
            display_name = None
            
            if user_profile and 'display_name' in user_profile:
                display_name = user_profile['display_name']
            
            # Save to session
            session['user_id'] = user_id
            session['user_email'] = user_email
            session['user_token'] = user_token
            session['display_name'] = display_name
            
            # Create user object and login
            user_obj = User(user_id, user_email, user_token, display_name)
            login_user(user_obj)
            
            # Redirect to home
            flash("Login successful!", "success")
            return redirect(url_for('index'))
            
        except Exception as e:
            logger.error(f"Login error: {str(e)}")
            error = "Invalid email or password. Please try again."
    
    return render_template('login.html', error=error, register=False)

@app.route('/register', methods=['GET', 'POST'])
def register():
    """Handle registration with email/password."""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    error = None
    success = None
    
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        try:
            # Create user with email/password
            user = auth.create_user_with_email_and_password(email, password)
            
            success = "Registration successful! You can now log in."
            
        except Exception as e:
            logger.error(f"Registration error: {str(e)}")
            error = "Registration failed. Please try again with a different email or stronger password."
    
    return render_template('login.html', error=error, success=success, register=True)

@app.route('/reset-password', methods=['GET', 'POST'])
def reset_password():
    """Handle password reset request."""
    error = None
    success = None
    
    if request.method == 'POST':
        email = request.form.get('email')
        try:
            auth.send_password_reset_email(email)
            success = "Password reset email sent. Check your inbox for instructions."
        except Exception as e:
            logger.error(f"Password reset error: {str(e)}")
            error = "Error sending password reset email. Please try again."
    
    return render_template('reset_password.html', error=error, success=success)
    
@app.route('/logout')
@login_required
def logout():
    """Handle logout."""
    logout_user()
    session.clear()
    flash("You have been logged out.", "info")
    return redirect(url_for('index'))

@app.route('/add-memory', methods=['GET', 'POST'])
@login_required
def add_memory():
    """Handle adding a new memory."""
    error = None
    
    if request.method == 'POST':
        try:
            # Get form data
            title = request.form.get('title')
            content = request.form.get('content')
            date = request.form.get('date')
            time = request.form.get('time')
            media_type = request.form.get('media_type')
            
            # Initialize media path
            media_path = None
            
            # Check which media was uploaded
            if media_type == 'image' and 'image' in request.files:
                file = request.files['image']
                if file and file.filename:
                    # Handle image upload
                    filename = secure_filename(file.filename)
                    file_path = f"images/{current_user.id}/{str(uuid.uuid4())}_{filename}"
                    storage.child(file_path).put(file)
                    media_path = storage.child(file_path).get_url()
            
            elif media_type == 'video' and 'video' in request.files:
                file = request.files['video']
                if file and file.filename:
                    # Handle video upload
                    filename = secure_filename(file.filename)
                    file_path = f"videos/{current_user.id}/{str(uuid.uuid4())}_{filename}"
                    storage.child(file_path).put(file)
                    media_path = storage.child(file_path).get_url()
            
            elif media_type == 'audio' and 'audio' in request.files:
                file = request.files['audio']
                if file and file.filename:
                    # Handle audio upload
                    filename = secure_filename(file.filename)
                    file_path = f"audios/{current_user.id}/{str(uuid.uuid4())}_{filename}"
                    storage.child(file_path).put(file)
                    media_path = storage.child(file_path).get_url()
            
            # Create memory object
            memory = {
                "title": title,
                "content": content,
                "date": date,
                "time": time,
                "created_at": datetime.now().isoformat(),
                "media": media_path,
                "media_type": media_type if media_path else None
            }
            
            # Save to Firebase
            db.child("memories").child(current_user.id).push(memory, token=current_user.token)
            
            flash("Memory saved successfully!", "success")
            return redirect(url_for('index'))
            
        except Exception as e:
            logger.error(f"Error adding memory: {str(e)}")
            error = f"Error saving memory: {str(e)}"
    
    return render_template('add_memory.html', error=error)

@app.route('/memory/<memory_id>')
@login_required
def view_memory(memory_id):
    """View a single memory."""
    try:
        # Get memory from Firebase
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if memory_ref.val():
            memory = memory_ref.val()
            memory['id'] = memory_id
            
            # Convert created_at string to datetime
            created_at = datetime.fromisoformat(memory['created_at'])
            memory['created_at'] = created_at
            
            return render_template('view_memory.html', memory=memory)
        
        flash("Memory not found.", "danger")
        return redirect(url_for('index'))
        
    except Exception as e:
        logger.error(f"Error viewing memory: {str(e)}")
        flash(f"Error loading memory: {str(e)}", "danger")
        return redirect(url_for('index'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    """User profile page."""
    error = None
    success = None
    display_name = current_user.display_name
    
    if request.method == 'POST':
        try:
            # Get form data
            new_display_name = request.form.get('display_name')
            
            if not new_display_name or len(new_display_name.strip()) == 0:
                error = "Display name cannot be empty!"
            else:
                # Update user profile in Firebase
                db.child("users").child(current_user.id).update({
                    "display_name": new_display_name
                }, token=current_user.token)
                
                # Update session
                session['display_name'] = new_display_name
                
                # Update current user object
                current_user.display_name = new_display_name
                
                success = "Profile updated successfully!"
                display_name = new_display_name
        except Exception as e:
            error = f"Error updating profile: {str(e)}"
            logger.error(f"Error updating profile: {str(e)}")
    
    return render_template('profile.html', error=error, success=success, display_name=display_name)

@app.route('/memory/<memory_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_memory(memory_id):
    """Edit a memory."""
    error = None
    memory = None
    
    try:
        # Get memory from Firebase
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if not memory_ref.val():
            flash("Memory not found.", "danger")
            return redirect(url_for('index'))
        
        memory = memory_ref.val()
        memory['id'] = memory_id
        
        if request.method == 'POST':
            # Get form data
            title = request.form.get('title')
            content = request.form.get('content')
            date = request.form.get('date')
            time = request.form.get('time')
            
            # Update memory data
            memory_data = {
                'title': title,
                'content': content,
                'date': date,
                'time': time,
                'updated_at': datetime.now().isoformat()
            }
            
            # Keep existing media data
            if 'media' in memory:
                memory_data['media'] = memory['media']
            if 'media_type' in memory:
                memory_data['media_type'] = memory['media_type']
            if 'created_at' in memory:
                memory_data['created_at'] = memory['created_at']
            
            # Update in Firebase
            db.child("memories").child(current_user.id).child(memory_id).update(memory_data, token=current_user.token)
            
            flash("Memory updated successfully!", "success")
            return redirect(url_for('view_memory', memory_id=memory_id))
        
        return render_template('edit_memory.html', memory=memory, error=error)
        
    except Exception as e:
        logger.error(f"Error editing memory: {str(e)}")
        flash(f"Error editing memory: {str(e)}", "danger")
        return redirect(url_for('index'))

@app.route('/memory/<memory_id>/delete', methods=['POST'])
@login_required
def delete_memory(memory_id):
    """Delete a memory."""
    try:
        # Get memory first to check if it has media
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if memory_ref.val() and memory_ref.val().get('media'):
            # Delete media file from storage if exists
            # This would require extracting the file path from the URL, which can be complex
            # For simplicity, we're just deleting the database entry
            pass
        
        # Delete memory from database
        db.child("memories").child(current_user.id).child(memory_id).remove(token=current_user.token)
        
        flash("Memory deleted successfully!", "success")
        
    except Exception as e:
        logger.error(f"Error deleting memory: {str(e)}")
        flash(f"Error deleting memory: {str(e)}", "danger")
    
    return redirect(url_for('index'))

# Export routes
@app.route('/export/pdf/<memory_id>')
@login_required
def export_memory_pdf(memory_id):
    """Export a single memory as PDF."""
    try:
        # Get memory from Firebase
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if not memory_ref.val():
            flash("Memory not found.", "danger")
            return redirect(url_for('index'))
        
        memory = memory_ref.val()
        memory['id'] = memory_id
        
        # Create HTML content for WeasyPrint
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{memory['title']}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #3498db; }}
                .memory-date {{ color: #666; font-size: 14px; margin-bottom: 20px; }}
                .memory-content {{ white-space: pre-wrap; line-height: 1.5; }}
                img {{ max-width: 80%; height: auto; margin: 15px 0; display: block; }}
            </style>
        </head>
        <body>
            <h1>{memory['title']}</h1>
            <div class="memory-date">Date: {memory['date']} | Time: {memory['time']}</div>
            <div class="memory-content">{memory['content']}</div>
        """
        
        # Add media if it exists
        if 'media' in memory and memory.get('media_type') == 'image':
            html_content += f'<img src="{memory["media"]}" alt="Memory Image">'
        
        html_content += """
        </body>
        </html>
        """
        
        # Generate PDF
        pdf = HTML(string=html_content).write_pdf()
        
        # Create an in-memory file-like object
        pdf_io = io.BytesIO(pdf)
        pdf_io.seek(0)
        
        # Generate a safe filename
        safe_title = secure_filename(memory['title'])
        filename = f"memory_{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Send the PDF as a download
        return send_file(
            pdf_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        logger.error(f"Error exporting memory as PDF: {str(e)}")
        flash(f"Error exporting memory: {str(e)}", "danger")
        return redirect(url_for('view_memory', memory_id=memory_id))

@app.route('/export/docx/<memory_id>')
@login_required
def export_memory_docx(memory_id):
    """Export a single memory as DOCX."""
    try:
        # Get memory from Firebase
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if not memory_ref.val():
            flash("Memory not found.", "danger")
            return redirect(url_for('index'))
        
        memory = memory_ref.val()
        memory['id'] = memory_id
        
        # Create document
        doc = Document()
        doc.add_heading(memory['title'], 0)
        
        # Add date and time
        doc.add_paragraph(f"Date: {memory['date']} | Time: {memory['time']}")
        
        # Add content
        doc.add_paragraph(memory['content'])
        
        # Add media if it exists and is an image
        if 'media' in memory and memory.get('media_type') == 'image':
            # For images, we can add them to the document
            # This would require downloading the image, which can be complex
            # For simplicity, we'll just note that there was an image
            doc.add_paragraph("This memory contains an image that can be viewed in the web application.")
        
        # Save to a BytesIO object
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # Generate a safe filename
        safe_title = secure_filename(memory['title'])
        filename = f"memory_{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Send the DOCX as a download
        return send_file(
            docx_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error exporting memory as DOCX: {str(e)}")
        flash(f"Error exporting memory: {str(e)}", "danger")
        return redirect(url_for('view_memory', memory_id=memory_id))

@app.route('/export/json/<memory_id>')
@login_required
def export_memory_json(memory_id):
    """Export a single memory as JSON."""
    try:
        # Get memory from Firebase
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if not memory_ref.val():
            flash("Memory not found.", "danger")
            return redirect(url_for('index'))
        
        memory = memory_ref.val()
        memory['id'] = memory_id
        
        # Convert to JSON
        memory_json = json.dumps(memory, indent=4)
        
        # Create a BytesIO object
        json_io = io.BytesIO()
        json_io.write(memory_json.encode('utf-8'))
        json_io.seek(0)
        
        # Generate a safe filename
        safe_title = secure_filename(memory['title'])
        filename = f"memory_{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        # Send the JSON as a download
        return send_file(
            json_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/json'
        )
        
    except Exception as e:
        logger.error(f"Error exporting memory as JSON: {str(e)}")
        flash(f"Error exporting memory: {str(e)}", "danger")
        return redirect(url_for('view_memory', memory_id=memory_id))

@app.route('/export/xlsx/<memory_id>')
@login_required
def export_memory_xlsx(memory_id):
    """Export a single memory as XLSX."""
    try:
        # Get memory from Firebase
        memory_ref = db.child("memories").child(current_user.id).child(memory_id).get(token=current_user.token)
        
        if not memory_ref.val():
            flash("Memory not found.", "danger")
            return redirect(url_for('index'))
        
        memory = memory_ref.val()
        memory['id'] = memory_id
        
        # Create workbook and select active worksheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Memory"
        
        # Add headers
        headers = ["Title", "Date", "Time", "Content", "Has Media"]
        for col_num, header in enumerate(headers, 1):
            ws.cell(row=1, column=col_num, value=header)
        
        # Add memory data
        ws.cell(row=2, column=1, value=memory['title'])
        ws.cell(row=2, column=2, value=memory['date'])
        ws.cell(row=2, column=3, value=memory['time'])
        ws.cell(row=2, column=4, value=memory['content'])
        ws.cell(row=2, column=5, value='Yes' if 'media' in memory else 'No')
        
        # Format the cells
        for col_num in range(1, len(headers) + 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_num)].width = 20
        
        # Save to a BytesIO object
        xlsx_io = io.BytesIO()
        wb.save(xlsx_io)
        xlsx_io.seek(0)
        
        # Generate a safe filename
        safe_title = secure_filename(memory['title'])
        filename = f"memory_{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        # Send the XLSX as a download
        return send_file(
            xlsx_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        logger.error(f"Error exporting memory as XLSX: {str(e)}")
        flash(f"Error exporting memory: {str(e)}", "danger")
        return redirect(url_for('view_memory', memory_id=memory_id))

# Export all memories routes
@app.route('/export/all/pdf')
@login_required
def export_all_memories_pdf():
    """Export all memories as PDF."""
    try:
        # Get all memories from Firebase
        memories_ref = db.child("memories").child(current_user.id).get(token=current_user.token)
        
        if not memories_ref.val():
            flash("No memories found to export.", "warning")
            return redirect(url_for('index'))
        
        # Create HTML content for WeasyPrint
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Memory Diary - All Memories</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #3498db; text-align: center; margin-bottom: 30px; }}
                h2 {{ color: #2c3e50; margin-top: 40px; border-bottom: 1px solid #eee; padding-bottom: 10px; }}
                .memory-date {{ color: #666; font-size: 14px; margin-bottom: 10px; }}
                .memory-content {{ white-space: pre-wrap; line-height: 1.5; margin-bottom: 20px; }}
                img {{ max-width: 70%; height: auto; margin: 15px auto; display: block; }}
                .memory-container {{ border-bottom: 2px dashed #ccc; padding-bottom: 30px; margin-bottom: 30px; }}
                .footer {{ text-align: center; font-size: 12px; color: #777; margin-top: 50px; }}
            </style>
        </head>
        <body>
            <h1>Memory Diary - All Memories</h1>
            <p style="text-align: center;">Exported on {datetime.now().strftime('%B %d, %Y at %H:%M')}</p>
        """
        
        # Add each memory
        memories = []
        for key, value in memories_ref.val().items():
            memory = value
            memory['id'] = key
            memories.append(memory)
        
        # Sort memories by date (newest first)
        memories.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        for memory in memories:
            html_content += f"""
            <div class="memory-container">
                <h2>{memory['title']}</h2>
                <div class="memory-date">Date: {memory['date']} | Time: {memory['time']}</div>
                <div class="memory-content">{memory['content']}</div>
            """
            
            # Add media if it exists
            if 'media' in memory and memory.get('media_type') == 'image':
                html_content += f'<img src="{memory["media"]}" alt="Memory Image">'
            
            html_content += "</div>"
        
        html_content += """
            <div class="footer">Generated by Memory Diary App</div>
        </body>
        </html>
        """
        
        # Generate PDF
        pdf = HTML(string=html_content).write_pdf()
        
        # Create an in-memory file-like object
        pdf_io = io.BytesIO(pdf)
        pdf_io.seek(0)
        
        # Generate a filename
        filename = f"all_memories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Send the PDF as a download
        return send_file(
            pdf_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        logger.error(f"Error exporting all memories as PDF: {str(e)}")
        flash(f"Error exporting memories: {str(e)}", "danger")
        return redirect(url_for('index'))

@app.route('/export/all/docx')
@login_required
def export_all_memories_docx():
    """Export all memories as DOCX."""
    try:
        # Get all memories from Firebase
        memories_ref = db.child("memories").child(current_user.id).get(token=current_user.token)
        
        if not memories_ref.val():
            flash("No memories found to export.", "warning")
            return redirect(url_for('index'))
        
        # Create document
        doc = Document()
        doc.add_heading('Memory Diary - All Memories', 0)
        
        # Add export date
        doc.add_paragraph(f"Exported on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        
        # Sort memories by date (newest first)
        memories = []
        for key, value in memories_ref.val().items():
            memory = value
            memory['id'] = key
            memories.append(memory)
        
        memories.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        # Add each memory
        for memory in memories:
            doc.add_heading(memory['title'], level=1)
            doc.add_paragraph(f"Date: {memory['date']} | Time: {memory['time']}")
            
            # Add content
            doc.add_paragraph(memory['content'])
            
            # Add media note if exists
            if 'media' in memory and memory.get('media_type'):
                doc.add_paragraph(f"This memory contains a {memory.get('media_type')} attachment that can be viewed in the web application.")
            
            # Add separator
            doc.add_paragraph("---")
        
        # Save to a BytesIO object
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # Generate a filename
        filename = f"all_memories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Send the DOCX as a download
        return send_file(
            docx_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error exporting all memories as DOCX: {str(e)}")
        flash(f"Error exporting memories: {str(e)}", "danger")
        return redirect(url_for('index'))

@app.route('/export/all/json')
@login_required
def export_all_memories_json():
    """Export all memories as JSON."""
    try:
        # Get all memories from Firebase
        memories_ref = db.child("memories").child(current_user.id).get(token=current_user.token)
        
        if not memories_ref.val():
            flash("No memories found to export.", "warning")
            return redirect(url_for('index'))
        
        # Prepare data structure
        export_data = {
            "memories": [],
            "exported_at": datetime.now().isoformat(),
            "user_email": current_user.email
        }
        
        # Get all memories
        for key, value in memories_ref.val().items():
            memory = value
            memory['id'] = key
            export_data["memories"].append(memory)
        
        # Convert to JSON
        memories_json = json.dumps(export_data, indent=4)
        
        # Create a BytesIO object
        json_io = io.BytesIO()
        json_io.write(memories_json.encode('utf-8'))
        json_io.seek(0)
        
        # Generate a filename
        filename = f"all_memories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        # Send the JSON as a download
        return send_file(
            json_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/json'
        )
        
    except Exception as e:
        logger.error(f"Error exporting all memories as JSON: {str(e)}")
        flash(f"Error exporting memories: {str(e)}", "danger")
        return redirect(url_for('index'))

@app.route('/export/all/xlsx')
@login_required
def export_all_memories_xlsx():
    """Export all memories as XLSX."""
    try:
        # Get all memories from Firebase
        memories_ref = db.child("memories").child(current_user.id).get(token=current_user.token)
        
        if not memories_ref.val():
            flash("No memories found to export.", "warning")
            return redirect(url_for('index'))
        
        # Create workbook and select active worksheet
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "All Memories"
        
        # Add export info
        ws.cell(row=1, column=1, value="Memory Diary - All Memories Export")
        ws.cell(row=2, column=1, value=f"Exported on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        ws.cell(row=3, column=1, value=f"User: {current_user.email}")
        
        # Add headers at row 5
        headers = ["ID", "Title", "Date", "Time", "Content", "Has Media", "Media Type", "Created At"]
        for col_num, header in enumerate(headers, 1):
            ws.cell(row=5, column=col_num, value=header)
        
        # Get all memories and sort by date (newest first)
        memories = []
        for key, value in memories_ref.val().items():
            memory = value
            memory['id'] = key
            memories.append(memory)
        
        memories.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        # Add memory data starting at row 6
        for row_num, memory in enumerate(memories, 6):
            ws.cell(row=row_num, column=1, value=memory.get('id', ''))
            ws.cell(row=row_num, column=2, value=memory.get('title', ''))
            ws.cell(row=row_num, column=3, value=memory.get('date', ''))
            ws.cell(row=row_num, column=4, value=memory.get('time', ''))
            ws.cell(row=row_num, column=5, value=memory.get('content', ''))
            ws.cell(row=row_num, column=6, value='Yes' if 'media' in memory else 'No')
            ws.cell(row=row_num, column=7, value=memory.get('media_type', ''))
            ws.cell(row=row_num, column=8, value=memory.get('created_at', ''))
        
        # Format the cells
        for col_num in range(1, len(headers) + 1):
            col_letter = openpyxl.utils.get_column_letter(col_num)
            # Adjust width based on column content
            if col_num == 5:  # Content column
                ws.column_dimensions[col_letter].width = 50
            else:
                ws.column_dimensions[col_letter].width = 18
        
        # Save to a BytesIO object
        xlsx_io = io.BytesIO()
        wb.save(xlsx_io)
        xlsx_io.seek(0)
        
        # Generate a filename
        filename = f"all_memories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        # Send the XLSX as a download
        return send_file(
            xlsx_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        logger.error(f"Error exporting all memories as XLSX: {str(e)}")
        flash(f"Error exporting memories: {str(e)}", "danger")
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)